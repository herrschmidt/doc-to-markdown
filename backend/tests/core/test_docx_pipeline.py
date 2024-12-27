import os
import pytest
import logging
from docx import Document
from docling.document_converter import DocumentConverter, WordFormatOption
from docling.datamodel.base_models import InputFormat

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

@pytest.fixture
def test_docx_path(tmp_path):
    """Create a test DOCX file with various formatting."""
    doc = Document()
    
    # Add heading
    doc.add_heading('Test Document', 0)
    
    # Add normal paragraph
    doc.add_paragraph('This is a normal paragraph.')
    
    # Add bold and italic text
    p = doc.add_paragraph()
    p.add_run('This text is ').bold = False
    p.add_run('bold').bold = True
    p.add_run(' and this is ').bold = False
    p.add_run('italic').italic = True
    p.add_run('.')
    
    # Add bullet list
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    
    # Add numbered list
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Header 1'
    cells[1].text = 'Header 2'
    cells = table.rows[1].cells
    cells[0].text = 'Cell 1'
    cells[1].text = 'Cell 2'
    
    # Save the document
    docx_path = tmp_path / "test_document.docx"
    doc.save(str(docx_path))
    return str(docx_path)

@pytest.fixture
def converter():
    """Create a DocumentConverter instance with WordFormatOption."""
    return DocumentConverter(
        allowed_formats=[InputFormat.DOCX],
        format_options={
            InputFormat.DOCX: WordFormatOption()
        }
    )

def test_docx_basic_conversion(converter, test_docx_path):
    """Test basic DOCX conversion with headings and paragraphs."""
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    assert '# Test Document' in markdown
    assert 'This is a normal paragraph.' in markdown

def test_docx_formatting(converter, test_docx_path):
    """Test that bold and italic formatting is preserved."""
    logger.info(f"Testing DOCX formatting with file: {test_docx_path}")
    logger.info(f"DOCX file exists: {os.path.exists(test_docx_path)}")
    logger.info(f"DOCX file size: {os.path.getsize(test_docx_path)} bytes")
    
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    logger.info("Generated markdown content:")
    logger.info(markdown)
    logger.info("\nChecking for bold and italic formatting...")
    
    # Check for bold text
    if '**bold**' not in markdown:
        logger.error("Could not find '**bold**' in markdown output")
        logger.info("Found text fragments containing 'bold':")
        for line in markdown.split('\n'):
            if 'bold' in line:
                logger.info(f"  {line}")
    
    # Check for italic text
    if '*italic*' not in markdown:
        logger.error("Could not find '*italic*' in markdown output")
        logger.info("Found text fragments containing 'italic':")
        for line in markdown.split('\n'):
            if 'italic' in line:
                logger.info(f"  {line}")
    
    assert '**bold**' in markdown
    assert '*italic*' in markdown

def test_docx_lists(converter, test_docx_path):
    """Test that bullet and numbered lists are correctly converted."""
    logger.info(f"Testing DOCX lists with file: {test_docx_path}")
    
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    logger.info("Generated markdown content:")
    logger.info(markdown)
    logger.info("\nChecking for bullet points and numbered lists...")
    
    # Check bullet points
    if '* First bullet point' not in markdown:
        logger.error("Could not find '* First bullet point' in markdown output")
        logger.info("Found text fragments containing 'First bullet point':")
        for line in markdown.split('\n'):
            if 'First bullet point' in line:
                logger.info(f"  {line}")
    
    if '* Second bullet point' not in markdown:
        logger.error("Could not find '* Second bullet point' in markdown output")
        logger.info("Found text fragments containing 'Second bullet point':")
        for line in markdown.split('\n'):
            if 'Second bullet point' in line:
                logger.info(f"  {line}")
    
    # Check numbered list
    if '1. First numbered item' not in markdown:
        logger.error("Could not find '1. First numbered item' in markdown output")
        logger.info("Found text fragments containing 'First numbered item':")
        for line in markdown.split('\n'):
            if 'First numbered item' in line:
                logger.info(f"  {line}")
    
    if '2. Second numbered item' not in markdown:
        logger.error("Could not find '2. Second numbered item' in markdown output")
        logger.info("Found text fragments containing 'Second numbered item':")
        for line in markdown.split('\n'):
            if 'Second numbered item' in line:
                logger.info(f"  {line}")
    
    assert '* First bullet point' in markdown
    assert '* Second bullet point' in markdown
    assert '1. First numbered item' in markdown
    assert '2. Second numbered item' in markdown

def test_docx_tables(converter, test_docx_path):
    """Test that tables are correctly converted to markdown format."""
    logger.info(f"Testing DOCX tables with file: {test_docx_path}")
    
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    logger.info("Generated markdown content:")
    logger.info(markdown)
    logger.info("\nChecking for table formatting...")
    
    # Check table headers and separator
    if '| Header 1 | Header 2 |' not in markdown:
        logger.error("Could not find '| Header 1 | Header 2 |' in markdown output")
        logger.info("Found text fragments containing 'Header':")
        for line in markdown.split('\n'):
            if 'Header' in line:
                logger.info(f"  {line}")
    
    if '|---|---|' not in markdown:
        logger.error("Could not find table separator '|---|---|' in markdown output")
        logger.info("Found lines containing '|':")
        for line in markdown.split('\n'):
            if '|' in line:
                logger.info(f"  {line}")
    
    if '| Cell 1 | Cell 2 |' not in markdown:
        logger.error("Could not find '| Cell 1 | Cell 2 |' in markdown output")
        logger.info("Found text fragments containing 'Cell':")
        for line in markdown.split('\n'):
            if 'Cell' in line:
                logger.info(f"  {line}")
    
    assert '| Header 1 | Header 2 |' in markdown
    assert '|---|---|' in markdown
    assert '| Cell 1 | Cell 2 |' in markdown

def test_docx_invalid_file(converter, tmp_path):
    """Test handling of invalid DOCX files."""
    invalid_path = tmp_path / "invalid.docx"
    with open(invalid_path, 'w') as f:
        f.write('Not a valid DOCX file')
    
    with pytest.raises(Exception):
        converter.convert(str(invalid_path))

def test_docx_empty_file(converter, tmp_path):
    """Test conversion of an empty DOCX file."""
    doc = Document()
    empty_path = tmp_path / "empty.docx"
    doc.save(str(empty_path))
    
    result = converter.convert(str(empty_path))
    markdown = result.document.export_to_markdown()
    
    assert markdown.strip() == ''