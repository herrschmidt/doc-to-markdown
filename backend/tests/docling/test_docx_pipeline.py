import os
import pytest
import logging
from docx import Document
from docling.document_converter import DocumentConverter, WordFormatOption
from docling.datamodel.base_models import InputFormat

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

@pytest.fixture
def test_docx_path(tmp_path):
    """Create a test DOCX file with various formatting."""
    doc = Document()
    
    # Add heading
    doc.add_heading('Test Document', 0)
    
    # Add normal paragraph
    doc.add_paragraph('This is a normal paragraph.')
    
    # Add bold and italic text
    p = doc.add_paragraph()
    p.add_run('This text is ').bold = False
    p.add_run('bold').bold = True
    p.add_run(' and this is ').bold = False
    p.add_run('italic').italic = True
    p.add_run('.')
    
    # Add bullet list
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    
    # Add numbered list
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Header 1'
    cells[1].text = 'Header 2'
    cells = table.rows[1].cells
    cells[0].text = 'Cell 1'
    cells[1].text = 'Cell 2'
    
    # Save the document
    docx_path = tmp_path / "test_document.docx"
    doc.save(str(docx_path))
    return str(docx_path)

@pytest.fixture
def converter():
    """Create a DocumentConverter instance with WordFormatOption."""
    return DocumentConverter(
        allowed_formats=[InputFormat.DOCX],
        format_options={
            InputFormat.DOCX: WordFormatOption()
        }
    )

def test_docx_basic_conversion(converter, test_docx_path):
    """Test basic DOCX conversion with headings and paragraphs."""
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    assert '# Test Document' in markdown
    assert 'This is a normal paragraph.' in markdown

def test_docx_text_content(converter, test_docx_path):
    """Test that all text content is preserved in the markdown output."""
    logger.info(f"Testing DOCX text content with file: {test_docx_path}")
    logger.info(f"DOCX file exists: {os.path.exists(test_docx_path)}")
    logger.info(f"DOCX file size: {os.path.getsize(test_docx_path)} bytes")
    
    result = converter.convert(test_docx_path)
    markdown = result.document.export_to_markdown()
    
    logger.info("Generated markdown content:")
    logger.info(markdown)
    logger.info("\nChecking for text content...")
    
    # Check for presence of all text content, regardless of formatting
    expected_texts = [
        "Test Document",
        "This is a normal paragraph",
        "This text is bold",
        "this is italic",
        "First bullet point",
        "Second bullet point",
        "First numbered item",
        "Second numbered item",
        "Header 1",
        "Header 2",
        "Cell 1",
        "Cell 2"
    ]
    
    for text in expected_texts:
        if text not in markdown:
            logger.error(f"Could not find '{text}' in markdown output")
            logger.info(f"Markdown content: {markdown}")
        assert text in markdown

def test_docx_invalid_file(converter, tmp_path):
    """Test handling of invalid DOCX files."""
    invalid_path = tmp_path / "invalid.docx"
    with open(invalid_path, 'w') as f:
        f.write('Not a valid DOCX file')
    
    with pytest.raises(Exception):
        converter.convert(str(invalid_path))

def test_docx_empty_file(converter, tmp_path):
    """Test conversion of an empty DOCX file."""
    doc = Document()
    empty_path = tmp_path / "empty.docx"
    doc.save(str(empty_path))
    
    result = converter.convert(str(empty_path))
    markdown = result.document.export_to_markdown()
    
    assert markdown.strip() == ''